from flask import Flask, request, jsonify
from google.cloud import storage
from google.api_core.exceptions import NotFound, Conflict
from datetime import timedelta, datetime
from vertexai.preview.language_models import TextGenerationModel
import vertexai
from docx import Document
from docx.shared import Inches
from flask_cors import CORS
import os
import logging
import traceback
import math
from collections import Counter
from google.cloud import bigquery
import pandas as pd
import re
import unicodedata
import calendar
import io


logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)

CORS(app, 
     origins=["https://accesa-client.vercel.app"],
     methods=["GET", "PUT", "POST", "OPTIONS"],
     allow_headers=["Content-Type", "Authorization"],
     supports_credentials=True
)

BUCKET_NAME = "accesa-data-gather"

vertexai.init(project="accesa-equipo3", location="southamerica-east1")


def wait_for_table_data(bq_client, table_id, timeout=15, interval=3):
    import time
    start = time.time()
    while time.time() - start < timeout:
        df = bq_client.query(f"SELECT COUNT(*) as total FROM `{table_id}`").to_dataframe()
        if df["total"][0] > 0:
            return True
        time.sleep(interval)
    return False

def convert_month_to_abbr(month_str):
    """Convert 'YEAR-MONTH' to 'MONTH-YEAR'"""
    year, month = map(int, month_str.split('-'))
    month_names = {
        1: "Ene", 2: "Feb", 3: "Mar", 4: "Abr", 5: "May", 6: "Jun",
        7: "Jul", 8: "Ago", 9: "Sep", 10: "Oct", 11: "Nov", 12: "Dic"
    }
    return f"{month_names.get(month, '---')}-{str(year)[2:]}"

def fetch_monthly_automatismo(table_name):
    bq = bigquery.Client()
    query = f"""
        SELECT 
        * FROM `accesa-equipo3.accesa_dataset.{table_name}`
        ;
    """
    return bq.query(query).to_dataframe()

def fetch_monthly_seisonce(table_name):
    bq = bigquery.Client()
    query = f"""
        SELECT * FROM `accesa-equipo3.accesa_dataset.{table_name}`
    """
    return bq.query(query).to_dataframe()

def fetch_monthly_congestion(table_name):
    bq = bigquery.Client()
    query = f"""
        SELECT 
        * FROM `accesa-equipo3.accesa_dataset.{table_name}` 
        ;
    """
    return bq.query(query).to_dataframe()

def fetch_monthly_reclamos(table_name):
    bq = bigquery.Client()
    query = f"""
        SELECT 
        * FROM `accesa-equipo3.accesa_dataset.{table_name}`
        ;
    """
    return bq.query(query).to_dataframe()

def fetch_monthly_incidencias(table_name):
    bq = bigquery.Client()
    query = f"""
        SELECT 
        * FROM `accesa-equipo3.accesa_dataset.{table_name}`
        ;
    """
    return bq.query(query).to_dataframe()

def fetch_monthly_habilidad(table_name):
    bq = bigquery.Client()
    query = f"""
        SELECT 
        * FROM `accesa-equipo3.accesa_dataset.{table_name}`
        ;
    """
    return bq.query(query).to_dataframe()

def fetch_monthly_skill(table_name):
    bq = bigquery.Client()
    query = f"""
        SELECT 
        * FROM `accesa-equipo3.accesa_dataset.{table_name}`
        ;
    """
    return bq.query(query).to_dataframe()

def fetch_monthly_roaming(table_name):
    bq = bigquery.Client()
    query = f"""
        SELECT 
            cantidad_de_mensajes_entrantes,
            cantidad_de_mensajes_salientes,
            total_de_mensajes,
            promedio_de_mensajes_por_interaccion
        FROM `accesa-equipo3.accesa_dataset.{table_name}`
    """
    return bq.query(query).to_dataframe()

def summarize_dataframe(df, context=""):
    model = TextGenerationModel.from_pretrained("text-bison@001")
    prompt = f"""
Actúa como un analista de datos senior. Resume los siguientes datos ({context}) en español con una narrativa profesional, señalando tendencias, cumplimiento del nivel de servicio, o valores atípicos.
 {df.describe(include='all').to_string()}
"""
    return model.predict(prompt=prompt, temperature=0.7, max_output_tokens=1024).text

def add_report_header(doc, month):
    doc.add_heading("Informe Mensual de Gestión", 0)
    doc.add_paragraph(f"Móvil – {month}", style="Subtitle")
    doc.add_paragraph(f"Fecha de generación: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph("")

def add_general_summary(doc):
    doc.add_heading("Resumen General", level=1)
    doc.add_paragraph(
        "Informe mensual de gestión de servicios de Accesa Contact Center para los servicios 0800 6611 y *611 de atención de clientes de Móvil Antel y 0800 2466, atención a Agentes de Venta y Clientes Internos. El servicio se atiende todos los días del año de 0 a 24 horas."
    )
    doc.add_paragraph("")
    
def add_header_image(doc, image_path):
    if os.path.exists(image_path):
        section = doc.sections[0]
        header = section.header
        paragraph = header.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(5.0)) 
    else:
        print(f"Header image not found: {image_path}")

def add_footer_image(doc, image_path):
    if os.path.exists(image_path):
        section = doc.sections[0]
        footer = section.footer
        paragraph = footer.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(6.0))  
    else:
        print(f"Footer image not found: {image_path}")
         
def add_summary_table_movil_611(doc, month_str, metrics):
    """
    Adds the custom summary table for Antel - Móvil 611 section with specific fields and structure.
    """
    doc.add_paragraph("")
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "ANTEL - MÓVIL 611"
    hdr_cells[1].text = month_str

    rows = [
        ("Llamadas al servicio", "llamadas_al_servicio", "int"),
        ("Llamadas atendidas totales", "llamadas_atendidas_totales", "int"),
        ("Llamadas abandonadas", "llamadas_abandonadas", "int"),
        ("% Llamadas no atendidas", "porcentaje_no_atendidas", "percent"),
        ("Cumplimiento Nivel de Servicio 80/20", "nivel_de_servicio_80_20", "percent"),
        ("Índice de respuesta", "indice_respuesta", "percent"),
        ("TRSAC", "trsac", "int"),
        ("Promedio operación (segundos)", "promedio_tiempo_operacion_segundos", "float"),
        ("Tiempo total atención (horas)", "tiempo_total_atencion_horas", "float"),
        ("Congestión", "congestion_6611", "percent")
    ]

    for label, key, dtype in rows:
        row = table.add_row().cells
        value = metrics.get(key)
        if value is None:
            display_value = "N/A"
        else:
            try:
                if dtype == "percent":
                    display_value = f"{float(value):.2f}%"
                elif dtype == "float":
                    display_value = f"{float(value):.2f}"
                else:
                    display_value = str(int(value))
            except (ValueError, TypeError):
                display_value = "N/A"

        row[0].text = label
        row[1].text = display_value

    promedio = metrics.get("promedio_llamadas_por_dia")
    if promedio is not None:
        try:
            promedio_val = int(promedio)
            doc.add_paragraph("")
            p = doc.add_paragraph()
    
            p.add_run("El promedio de llamadas diarias ingresado al servicio en el mes fue de ")
            run_promedio = p.add_run(f"{promedio_val}")
            run_promedio.bold = True
            
            doc.add_paragraph("")
        except (ValueError, TypeError):
            pass


def add_summary_table_reclamos(doc, month_str, metrics):
    """
    Adds a custom summary table for Reclamos section.
    """
    doc.add_paragraph("")  
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Mes"
    hdr_cells[1].text = "Campaña"
    hdr_cells[2].text = "Total llamadas"
    hdr_cells[3].text = "Tiempo total"

    val_cells = table.rows[1].cells
    val_cells[0].text = month_str
    val_cells[1].text = "Reclamos_611"
    val_cells[2].text = str(metrics.get("total_llamadas", "N/A"))
    val_cells[3].text = str(metrics.get("total_tiempo_llamadas", "N/A"))     

def add_reclamos_motivos_table(doc, metrics):
    motivo_map = metrics.get("motivo_manejo_map", {})

    motivo_map = {
        motivo: cantidad
        for motivo, cantidad in motivo_map.items()
        if motivo.lower().strip() != "inin-wrap-up-timeout"
    }

    if not motivo_map:
        return

    sorted_motivos = sorted(motivo_map.items(), key=lambda x: x[1], reverse=True)

    doc.add_paragraph().add_run("Motivos IZI 611").bold = True

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    header_cells[0].text = "Motivos IZI 611"
    header_cells[1].text = "Cantidad"

    total_count = 0

    for motivo, cantidad in sorted_motivos:
        row_cells = table.add_row().cells
        row_cells[0].text = motivo
        row_cells[1].text = str(int(cantidad))
        total_count += cantidad

    total_row = table.add_row().cells
    total_row[0].text = "Total"
    total_row[0].paragraphs[0].runs[0].bold = True
    total_row[1].text = f"{int(total_count):,}".replace(",", ".")
    total_row[1].paragraphs[0].runs[0].bold = True

    doc.add_paragraph("")


def add_summary_table_seisonce(doc, metrics):
    """
    Agrega un título, párrafo descriptivo, y tabla de motivos válidos para 611.
    """
    doc.add_paragraph("") 
    doc.add_paragraph("Motivos de los contactos", style="Heading 2")

    total_motivos_validos = metrics.get("total_motivos_validos", 0)
    porcentaje_valido = metrics.get("porcentaje_valido", 0)

    p = doc.add_paragraph()
    
    p.add_run("Durante el mes se registraron ")
    run_motivos = p.add_run(f"{total_motivos_validos:,}")
    run_motivos.bold = True
    p.add_run(" motivos, lo que corresponde al ")
    
    run_porcentaje = p.add_run(f"{porcentaje_valido:.2f}%")
    run_porcentaje.bold = True
    p.add_run(" de las llamadas atendidas.")
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    header_cells = table.rows[0].cells
    header_cells[0].text = "Motivo"
    header_cells[1].text = "Cantidad"

    motivos = metrics.get("motivos", [])
    cantidades = metrics.get("cantidades", [])

    for motivo, cantidad in zip(motivos, cantidades):
        row_cells = table.add_row().cells
        row_cells[0].text = motivo.capitalize()
        row_cells[1].text = str(int(cantidad))

    doc.add_paragraph("")  


def add_summary_table_automatismos(doc, month_str, metrics):
    """
    Adds a 3-column automatismos table:
    [Tarjeta Móvil Agencias - <month>] | Cantidad | Porcentaje
    
    """
    doc.add_paragraph("") 
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    header_cells = table.rows[0].cells
    header_cells[0].text = f"Tarjeta Móvil Agencias - {month_str}"
    header_cells[1].text = "Cantidad"
    header_cells[2].text = "Porcentaje"

    def fmt(key, dtype):
        value = metrics.get(key)
        if value is None:
            return "N/A"
        try:
            return f"{float(value):.2f}%" if dtype == "percent" else f"{int(value):,}".replace(",", ".")
        except Exception:
            return "N/A"

  
    row = table.add_row().cells
    row[0].text = "Total éxito"
    row[1].text = fmt("total_correcto", "int")
    row[2].text = fmt("porcentaje_correcto", "percent")

    
    row = table.add_row().cells
    row[0].text = "Total errores"
    row[1].text = fmt("total_error", "int")
    row[2].text = fmt("porcentaje_error", "percent")

   
    row = table.add_row().cells
    row[0].text = "Total"
    row[1].text = fmt("total_automatismos", "int")
    row[2].text = "100.00%" 

    doc.add_paragraph("") 


def add_incidencias_bullet_section(doc, metrics):
    """
    Adds a heading and a bulleted list of incidencias, with their dates and descriptions.
    """
    fechas = metrics.get("fecha_incidente", [])
    descripciones = metrics.get("descripcion_incidente", [])

    doc.add_paragraph("")
    if fechas is None or descripciones is None or len(fechas) == 0:
        doc.add_paragraph("No se registraron incidencias que afectaran el servicio móvil este mes.", style="Subtitle")
        return

    for fecha, descripcion in zip(fechas, descripciones):
        item = f"{fecha}: {descripcion}"
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph("")

            
def build_report(month, metrics_by_section):
    doc = Document()
        
    add_header_image(doc, "assets/accesa-header-doc.png")
    add_footer_image(doc, "assets/accesa-footer-doc.png")
    
    add_report_header(doc, month)
    add_general_summary(doc)


    doc.add_heading("Indicadores de Gestión de las Llamadas", level=1)

    habilidad = metrics_by_section.get("habilidad_data", {})
    skill = metrics_by_section.get("skill_data", {})
    congestion = metrics_by_section.get("congestion_data", {})
    incidencias = metrics_by_section.get("incidencias_data", {})

    combined_611_metrics = {
        "llamadas_al_servicio": skill.get("llamadas_al_servicio"),
        "llamadas_atendidas_totales": skill.get("llamadas_atendidas_totales"),
        "llamadas_abandonadas": skill.get("llamadas_abandonadas"),
        "porcentaje_no_atendidas": habilidad.get("porcentaje_no_atendidas"),
        "nivel_de_servicio_80_20": skill.get("nivel_de_servicio_80_20"),
        "indice_respuesta": habilidad.get("indice_respuesta"),
        "trsac": skill.get("trsac"),
        "promedio_tiempo_operacion_segundos": skill.get("promedio_tiempo_operacion_segundos"),
        "tiempo_total_atencion_horas": skill.get("tiempo_total_atencion_horas"),
        "congestion_6611": congestion.get("congestion_6611"),
        "promedio_llamadas_por_dia": habilidad.get("promedio_llamadas_por_dia"),
    }

    add_summary_table_movil_611(doc, convert_month_to_abbr(month), combined_611_metrics)
    
    doc.add_paragraph("")
    if incidencias:
        p = doc.add_paragraph()
        p.add_run("Incidencias que afectaron el servicio móvil en el mes:").bold = True
        add_incidencias_bullet_section(doc, incidencias)
        
    p = doc.add_paragraph()
    p.add_run("Gestión Sistema Reclamos").bold = True

    doc.add_paragraph(
        "Además de la atención de la línea los Agentes del servicio Móvil realizan la gestión de la bandeja "
        "en el Sistema Reclamos de ANTEL. A dicha carpeta llegan las consultas de clientes provenientes de "
        "Whatsapp, la App MiAntel, de la Web MiAntel y de las Oficinas Comerciales."
    )
   
    if "reclamos_data" in metrics_by_section:
        add_summary_table_reclamos(
        doc, convert_month_to_abbr(month), metrics_by_section["reclamos_data"])
        add_reclamos_motivos_table(doc, metrics_by_section["reclamos_data"])
        doc.add_paragraph("")
    
    if "roaming_data" in metrics_by_section:
        doc.add_paragraph("")
        doc.add_paragraph().add_run("Asistencia por Roaming vía WhatsApp (092611611 opción 7)").bold = True
        doc.add_paragraph("")
        roaming_metrics = metrics_by_section["roaming_data"]

        table = doc.add_table(rows=5, cols=2) 
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].merge(hdr_cells[1])
        hdr_cells[0].text = convert_month_to_abbr(month)  
        hdr_cells[0].paragraphs[0].runs[0].bold = True

        rows = [
        ("Cantidad de mensajes entrantes", "mensajes_entrantes"),
        ("Cantidad de mensajes salientes", "mensajes_salientes"),
        ("Total de mensajes", "total_mensajes"),
        ("Promedio de mensajes por interacción", "promedio_mensajes_por_interaccion"),
    ]

        for idx, (label, key) in enumerate(rows, start=1):
            row_cells = table.rows[idx].cells
            row_cells[0].text = label
            value = roaming_metrics.get(key)
            row_cells[1].text = str(value if value is not None else "N/A")

    if "611_data" in metrics_by_section:
        add_summary_table_seisonce(
        doc, metrics_by_section["611_data"]
    )

      
    if "congestion_data" in metrics_by_section:
        doc.add_paragraph().add_run("Congestión").bold = True
        congestion_raw = metrics_by_section["congestion_data"].get("congestion_6611")

        try:
            congestion_value = f"{float(congestion_raw):.2f}%"
        except (ValueError, TypeError):
            congestion_value = "N/A"

        promedio_operacion = skill.get("promedio_tiempo_operacion_segundos", "N/A")
        trsac = skill.get("trsac", "N/A")

        doc.add_paragraph(f"La congestión del mes fue de: {congestion_value}")
        doc.add_paragraph(
        f"Durante el mes el tiempo de operación promedio fue de {promedio_operacion} segundos y el TRSAC fue de {trsac} segundos"
        )
   
    if "automatismo_data" in metrics_by_section:
        p = doc.add_paragraph()
        p.add_run("Automatismos").bold = True

        doc.add_paragraph(
            "El uso de automatismos libera recursos para la atención de otras consultas más complejas. "
            "Actualmente está operativa una automatización permanente para el servicio brindado a Móvil Antel."
        )
        add_summary_table_automatismos(
        doc, convert_month_to_abbr(month), metrics_by_section["automatismo_data"]
        )
        doc.add_paragraph("Las horas incurridas en los automatismos no se computan como horas de operación mensual.")

    filename = f"Informe Móvil_{month}.docx"
    doc.save(filename)
    return filename



@app.route("/generate-signed-url", methods=["POST", "OPTIONS"])
def get_signed_url():
    try:
        if request.method == "OPTIONS":
            response = jsonify({"status": "ok"})
            return response, 200
        
        if not request.is_json:
            logger.error("Request not json")
            return jsonify({"error": "Request must be json"}), 400
        
        data = request.get_json()
        if not data:
            logger.error("No data")
            return jsonify({"error": "No data provided"}), 400
            
        file_name = data.get("file_name")
        content_type = data.get("content_type", "application/octet-stream")

        if not file_name:
            logger.error("file_name not found")
            return jsonify({"error": "file_name is required"}), 400

        storage_client = storage.Client()
        bucket = storage_client.bucket(BUCKET_NAME)
        blob = bucket.blob(f"accesa-data/{file_name}")

        url = blob.generate_signed_url(
            version="v4",
            expiration=timedelta(minutes=15),
            method="PUT",
            content_type=content_type,
        )

        logger.info(f"Generated signed url for file: {file_name}")
        return jsonify({"signed_url": url}), 200
        
    except Exception as e:
        logger.error(f"Error get_signed_url: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({"error": "Internal server error"}), 500


@app.route("/pubsub-notify-hook", methods=["POST"])
def pubsub_hook():
    try:
        envelope = request.get_json()
        if not envelope:
            msg = "No message"
            logger.error(msg)
            return msg, 400

        pubsub_message = envelope.get("message")
        if not pubsub_message:
            msg = "Invalid Pub/Sub format"
            logger.error(msg)
            return msg, 400

        import base64
        data = base64.b64decode(pubsub_message["data"]).decode("utf-8")
        logger.info(f"Pub/Sub Message: {data}")


        import json
        event_data = json.loads(data)
        file_name = event_data["name"]
        if file_name.startswith("reports/") or file_name.endswith(".docx"):
            logger.info(f"Archivo de reporte : {file_name}")
            return "Archivo de reporte omitido", 200
        
        try:
            logger.info(f"Triggering file process for: {file_name}")
            process_file(file_name)
        except Exception as e:
            logger.error(f"Error processing file {file_name}: {str(e)}")

        return "OK", 200
    except Exception as e:
        logger.error(f"hook failed: {e}")
        logger.error(traceback.format_exc())
        return "Internal Error", 500

@app.route("/generate-report", methods=["POST"])
def generate_report():
    handler_map = {
    "roaming_data" : (fetch_monthly_roaming, calc_roaming),
    "habilidad_data": (fetch_monthly_habilidad, calc_habilidad),
    "611_data": (fetch_monthly_seisonce, calc_seisonce),
    "skill_data": (fetch_monthly_skill, calc_skill),
    "reclamos_data":(fetch_monthly_reclamos, calc_reclamos),
    "automatismo_data": (fetch_monthly_automatismo, calc_automatismo),
    "incidencias_data" : (fetch_monthly_incidencias, calc_incidencias),
    "congestion_data": (fetch_monthly_congestion, calc_congestion),
    }
    try:
        data = request.get_json()
        month = data.get("month")

        metrics_by_section = {}
        bq_client = bigquery.Client()

        for table, (fetch_fn, calc_fn) in handler_map.items():
            table_id = f"accesa-equipo3.accesa_dataset.{table}"
            if not wait_for_table_data(bq_client, table_id):
                logger.error(f"Tiempo agotado para la tabla: {table}")
                raise RuntimeError(f"Timeout for data in table: {table}")
            try:
                df = fetch_fn(table)
                if not df.empty:
                    metrics = calc_fn(df)
                    logger.info(f"calcs returned: {metrics}")
                    metrics_by_section[table] = metrics
            except Exception as e:
                logger.warning(f"Failed to fetch data for {table}: {e}")

        if not metrics_by_section:
            return jsonify({"error": f"No data for period {month}"}), 404

        report_filename = build_report(month, metrics_by_section)

        storage_client = storage.Client()
        bucket = storage_client.bucket(BUCKET_NAME)
        blob = bucket.blob(f"reports/report_{month}.docx")
        blob.upload_from_filename(report_filename)

        # Download URL
        url = blob.generate_signed_url(
            version="v4",
            expiration=timedelta(minutes=15),
            method="GET"
        )
        return jsonify({"download_url": url}), 200

    except Exception as e:
        logger.error(f"Failed to generate report: {e}")
        logger.error(traceback.format_exc())
        return jsonify({"error": "Internal server error"}), 500


@app.route("/health", methods=["GET"])
def health_check():
    return jsonify({"status": "healthy"}), 200


@app.errorhandler(404)
def not_found(error):
    return jsonify({"error": "Not found"}), 404


@app.errorhandler(500)
def internal_error(error):
    logger.error(f"Internal server error: {error}")
    return jsonify({"error": "Internal server error"}), 500

def clean_column_name(name):
    name = unicodedata.normalize('NFKD', name).encode('ASCII', 'ignore').decode('utf-8')

    name = name.strip().lower().replace(" ", "_")

    name = re.sub(r"[^\w]", "_", name)

    name = re.sub(r"_+", "_", name)

    if not re.match(r"^[a-zA-Z_]", name):
        name = f"col_{name}"
        
    return name[:128]

# Función para procesar los archivos y cargarlos a BigQuery, llamada por Pub/Sub
def process_file(file_name):
    try:
        if not file_name:
            raise ValueError("Missing file_name in process_file")

        storage_client = storage.Client()
        bucket = storage_client.bucket(BUCKET_NAME)
        blob = bucket.blob(file_name)
        file_ext = file_name.split(".")[-1]

        file_data = blob.download_as_bytes()
        
        if file_ext == "csv":
            df = pd.read_csv(io.BytesIO(file_data))
        elif file_ext in ["xls", "xlsx"]:
            df = pd.read_excel(io.BytesIO(file_data), header=None, dtype=str)
            logger.info(f"Dataframe actual: {df.shape}")
            if df.shape[1] == 1 and isinstance(df.iloc[0, 0], str) and df.iloc[0, 0].count(',') > 1:
                logger.warning("Improperly formatted Excel file with comma-delimited content.")

                split_df = df.iloc[:, 0].apply(lambda x: pd.Series(str(x).split('","')))

                split_df.columns = split_df.iloc[0].fillna("col_unnamed")
                df = split_df.iloc[1:].reset_index(drop=True)
            else:
                if "roaming" in file_name.lower():
                    df = pd.read_excel(io.BytesIO(file_data), skiprows=6)
                elif "skill" in file_name.lower():
                    df = pd.read_excel(io.BytesIO(file_data), skiprows=3)
                elif "automatismo" in file_name.lower():
                    df = pd.read_excel(io.BytesIO(file_data), skiprows=4)
                elif "habilidad" in file_name.lower():
                    df = pd.read_excel(io.BytesIO(file_data), skiprows=1)
                elif "congestion" in file_name.lower():
                    df = pd.read_excel(io.BytesIO(file_data), skiprows=4)
                else:
                    df = pd.read_excel(io.BytesIO(file_data))
            
        else:
            raise ValueError(f"error, unsupported file format {file_ext}")
        logger.info(f"Dataframe shape: {df.shape}")
        df.columns = [clean_column_name(col) for col in df.columns]
        logger.info(f"columnas procesadas: {df.columns.tolist()}")

        project_id = "accesa-equipo3"
        dataset_id = "accesa_dataset"
        file_name = file_name.lower()
        if "roaming" in file_name:
            table_name = "roaming_data"
        elif "automatismo" in file_name:
            table_name = "automatismo_data"
        elif "incidencias" in file_name:
            table_name = "incidencias_data"
        elif "congestion" in file_name:
            table_name = "congestion_data"
        elif "habilidad" in file_name:
            table_name = "habilidad_data"
        elif "skill" in file_name:
            table_name = "skill_data"
        elif "reclamos" in file_name:
            table_name = "reclamos_data"
        elif "611" in file_name:
            table_name = "611_data"
        else:
            table_name = "sin_categorizar"

        table_id = f"{project_id}.{dataset_id}.{table_name}"

        bq_client = bigquery.Client(project=project_id)

        dataset_ref = f"{project_id}.{dataset_id}"

        try:
            bq_client.get_dataset(dataset_ref)
            logger.info(f"Dataset {dataset_id} already exists.")
        except NotFound:
            try:
                dataset = bigquery.Dataset(dataset_ref)
                dataset.location = "southamerica-east1"
                bq_client.create_dataset(dataset)
                logger.info(f"Created dataset: {dataset.dataset_id}")
            except Conflict:
                logger.info(f"Dataset {dataset_ref} already exists (race condition).")
        job = bq_client.load_table_from_dataframe(df, table_id)
        job.result()

        logger.info(f"archivo: {file_name} cargado a: {table_name}"), 200

    except Exception as e:
        logger.error(f"Error in process_upload: {str(e)}")
        logger.error(traceback.format_exc())
        raise

def calc_habilidad(df):
    total_llamadas = df["ofrecidas"].sum()
    atendidas = df["atendidas"].sum()
    abandonadas = df["abandonadas"].sum()
    dias_mes = df["fecha"].nunique()

    porcentaje_no_atendidas = (abandonadas / total_llamadas) * 100 if total_llamadas else 0
    indice_respuesta = 100 - porcentaje_no_atendidas
    promedio_llamadas = total_llamadas / dias_mes if dias_mes else 0

    return {
        "llamadas_servicio": total_llamadas,
        "llamadas_atendidas": atendidas,
        "llamadas_abandonadas": abandonadas,
        "porcentaje_no_atendidas": round(porcentaje_no_atendidas, 2),
        "indice_respuesta": round(indice_respuesta, 2),
        "promedio_llamadas_por_dia": math.trunc(promedio_llamadas),
    }

def calc_skill(df):
    def parse_trsac_string(value):
        try:
            h, m, s = map(int, value.split(":"))
            return h * 3600 + m * 60 + s
        except:
            return 0

    trsac_value = df["demora_en_atender"][1] if len(df["demora_en_atender"]) > 1 else "00:00:00"
    trsac_antel_movil = parse_trsac_string(trsac_value)

    antel_movil_ofrecidas = df["ofrecidas"][1] if len(df["ofrecidas"]) > 1 else 0
    antel_movil_contestadas = df["contestadas"][1] if len(df["contestadas"]) > 1 else 0
    antel_movil_abandonadas = df["abandonadas"][1] if len(df["abandonadas"]) > 1 else 0
    referentes_movil_ofrecidas = df["ofrecidas"][0] if len(df["ofrecidas"]) > 0 else 0
    horas_operacion = df["horas_operacion"][1] if len(df["horas_operacion"]) > 1 else 0
    promedio_tiempo_operacion_segundos = df["tiempo_operacion"][1] if len(df["tiempo_operacion"]) > 1 else 0
    calidad = df["calidad_"][1] if len(df["calidad_"]) > 1 else 0
    nivel_de_servicio = round((calidad * 100) / 80, 2)

    return {
        "trsac": trsac_antel_movil,
        "llamadas_al_servicio": antel_movil_ofrecidas,
        "llamadas_atendidas_totales": antel_movil_contestadas,
        "llamadas_abandonadas" : antel_movil_abandonadas,
        "referentes_movil_ofrecidas": referentes_movil_ofrecidas,
        "tiempo_total_atencion_horas": horas_operacion,
        "nivel_de_servicio_80_20": nivel_de_servicio,
        "promedio_tiempo_operacion_segundos": promedio_tiempo_operacion_segundos,
    }


def calc_congestion(df):
    df_seis_uno = df["col_6611_"].sum()
    df_uno_dos_uno = df["col_121_"].sum()
    
    return {
        "congestion_6611": float(df_seis_uno)
    }
def calc_seisonce(df):
    df["acumulado_o_detallado"] = df["acumulado_o_detallado"].astype(str).str.strip().str.lower()

    df_acumulado = df[df["acumulado_o_detallado"] == "acumulado"]
    df_acumulado["manejo"] = pd.to_numeric(df_acumulado["manejo"], errors="coerce").fillna(0)

    excluded = {"inin-wrap-up-deleted", "inin-wrap-up-timeout", "default wrap-up code"}

    df_validos = df_acumulado[~df_acumulado["nombre_de_codigo_de_conclusion"].str.lower().isin(excluded)]

    total_motivos_validos = df_validos["manejo"].sum()
    total_motivos = df_acumulado["manejo"].sum()

    porcentaje_valido = (total_motivos_validos / total_motivos * 100) if total_motivos else 0

    motivos = df_validos["nombre_de_codigo_de_conclusion"].astype(str).tolist()
    cantidades = df_validos["manejo"].tolist()

    return {
        "total_motivos_validos": int(total_motivos_validos),
        "porcentaje_valido": round(porcentaje_valido, 2),
        "motivos": motivos,
        "cantidades": cantidades
    }


def calc_roaming(df):
    mensajes_entrantes = df["cantidad_de_mensajes_entrantes"].iloc(0)
    mensajes_salientes = df["cantidad_de_mensajes_salientes"].iloc(0)
    total_mensajes = mensajes_entrantes + mensajes_salientes
    promedio_mensajes_por_interaccion = df["promedio_de_mensajes_por_interaccion"].iloc(0)

    return {
        "mensajes_entrantes": mensajes_entrantes,
        "mensajes_salientes": mensajes_salientes,
        "total_mensajes": total_mensajes,
        "promedio_mensajes_por_interaccion": promedio_mensajes_por_interaccion,
    }
    

def calc_automatismo(df):
    last_row_correcto = df["total_correcto"].iloc[-1]
    last_row_error = df["total_error"].iloc[-1]

    total_correcto = df["total_correcto"].sum() - last_row_correcto
    total_error = df["total_error"].sum() - last_row_error

    porcentaje_correcto = ((total_correcto * 100 )/ (total_correcto + total_error)) if (total_correcto + total_error) > 0 else 0
    porcentaje_error = ((total_error * 100) / (total_correcto + total_error)) if (total_correcto + total_error) > 0 else 0
    total_automatismos = total_correcto + total_error

    return {
        "total_correcto": total_correcto,
        "total_error": total_error,
        "total_automatismos": total_automatismos,
        "porcentaje_correcto": round(porcentaje_correcto, 2),
        "porcentaje_error": round(porcentaje_error, 2),
    }

def ms_to_hms(ms):
    seconds = int(ms / 1000)
    hours = seconds // 3600
    minutes = (seconds % 3600) // 60
    seconds = seconds % 60
    return f"{hours:02}:{minutes:02}:{seconds:02}"

def calc_incidencias(df):
    df["responsabilidad"] = df["responsabilidad"].str.lower()
    df = df[df["responsabilidad"] == "cliente"]
    
    fecha_incidente = df["fecha"]
    descripcion_incidente = df["descripcion"]
    
    return {
        "fecha_incidente" : fecha_incidente,
        "descripcion_incidente" : descripcion_incidente
    }
    
    
def calc_reclamos(df):
    
    df["acumulado_o_detallado"] = (
        df["acumulado_o_detallado"]
        .astype(str)
        .str.strip()
        .str.strip('"')
        .str.lower()
    )

    logger.info(
        "Unique values in acumulado_o_detallado before filtering: %s",
        df["acumulado_o_detallado"].unique(),
    )
    logger.info("DF size before filtering: %d", len(df))

    df = df[df["acumulado_o_detallado"] == "detallado"]
    logger.info("DF size after filtering: %d", len(df))

    df["manejo_total"] = pd.to_numeric(
        df["manejo_total"].astype(str).str.strip().str.strip('"'),
        errors="coerce"
    ).fillna(0)

    df["manejo"] = pd.to_numeric(
        df["manejo"].astype(str).str.strip().str.strip('"'),
        errors="coerce"
    ).fillna(0)
    excluded_codes = {"inin-wrap-up-timeout", "inin-wrap-up-delete", "default wrap-up code"}
    motivo_manejo = df[~df["nombre_de_codigo_de_conclusion"].str.lower().isin(excluded_codes)]

    motivo_counts = Counter()
    for motivo, cantidad in zip(motivo_manejo["nombre_de_codigo_de_conclusion"], motivo_manejo["_manejo_"]):
        motivo_counts[motivo.strip()] += cantidad

    total_tiempo_llamadas = df["manejo_total"].sum()
    total_llamadas = df["manejo"].sum()

    logger.info("Total manejo: %s", total_llamadas)
    logger.info("Total manejo total (ms): %s", total_tiempo_llamadas)

    nombre_de_cola = df["nombre_de_cola"].dropna().astype(str).to_list()
    nombre_de_codigo_de_conclusion = df["nombre_de_codigo_de_conclusion"].dropna().astype(str).to_list()

    return {
        "total_tiempo_llamadas": ms_to_hms(total_tiempo_llamadas),
        "total_llamadas": int(total_llamadas),
        "nombre_de_cola": nombre_de_cola,
        "motivo_manejo_map": dict(motivo_counts),
        "nombre_de_codigo_de_conclusion": nombre_de_codigo_de_conclusion,
    }


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8080))
    logger.info(f"Starting server on port: {port}")
    app.run(host="0.0.0.0", port=port, debug=False)